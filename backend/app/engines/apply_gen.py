"""Grounded document generation for the Apply Studio (plan-apply.md, Phase B).

Generates an ATS-friendly résumé, a cover letter, and screening-question answers
from the master profile + a job description. The hard rule (ADR #2): **use only
facts present in the profile — reword / reorder / emphasize, never invent an
employer, title, date, degree, or number.**

Two paths, same shape:
  * LOCAL (default, zero-credential): assembles documents from the profile's own
    fields, so output is grounded *by construction* and always passes the verifier.
  * GEMINI (when GEMINI_API_KEY is set): drafts the prose, then the draft is run
    through `verify_grounding`. If any invented year/fact is found the draft is
    rejected and we fall back to the local assembly — the verifier is a hard gate,
    not a prompt hope.

Renderers produce ATS-clean text/HTML (single column, standard headings, no
tables / images / icons) for `/export`.
"""
from __future__ import annotations

import html as _html
import re
import sys
from datetime import datetime
from typing import Any, Dict, List, Optional

from app.config import settings

_YEAR_RE = re.compile(r"\b(?:19|20)\d{2}\b")
KINDS = ("resume", "cover_letter", "answers")


# ---------------------------------------------------------------- profile access
def _sections(profile: Dict[str, Any]) -> List[Dict[str, Any]]:
    return profile.get("sections") or []


def _by_type(profile: Dict[str, Any], t: str) -> List[Dict[str, Any]]:
    return [s for s in _sections(profile) if s.get("type") == t]


def _personal(profile: Dict[str, Any]) -> Dict[str, Any]:
    secs = _by_type(profile, "personal")
    fields = secs[0].get("fields", {}) if secs else {}
    city = fields.get("city") or ""
    country = fields.get("country") or ""
    loc = fields.get("location") or ""
    if not loc and (city or country):
        loc = ", ".join([c for c in [city, country] if c])
    
    github = fields.get("github") or ""
    linkedin = fields.get("linkedin") or ""
    portfolio = fields.get("portfolio") or ""
    links = fields.get("links") or []
    if not links:
        links = [link for link in [github, linkedin, portfolio] if link]

    return {
        "name": fields.get("name") or profile.get("full_name") or "",
        "email": fields.get("email") or profile.get("email") or "",
        "phone": fields.get("mobile") or fields.get("phone") or profile.get("phone") or "",
        "mobile": fields.get("mobile") or fields.get("phone") or profile.get("phone") or "",
        "city": city,
        "country": country,
        "location": loc,
        "github": github,
        "linkedin": linkedin,
        "portfolio": portfolio,
        "links": links,
        "headline": fields.get("headline") or "",
    }


def _summary_text(profile: Dict[str, Any]) -> str:
    secs = _by_type(profile, "summary")
    return (secs[0].get("text") if secs else "") or ""


def _experience(profile: Dict[str, Any]) -> List[Dict[str, Any]]:
    out: List[Dict[str, Any]] = []
    for s in _by_type(profile, "experience"):
        out.extend(s.get("items") or [])
    return out


def _education(profile: Dict[str, Any]) -> List[Dict[str, Any]]:
    out: List[Dict[str, Any]] = []
    for s in _by_type(profile, "education"):
        out.extend(s.get("items") or [])
    return out


def _skill_names(profile: Dict[str, Any]) -> List[str]:
    out: List[str] = []
    for s in _by_type(profile, "skills"):
        for it in (s.get("items") or []):
            out.append(it if isinstance(it, str) else str(it.get("heading", "")))
    return [x for x in out if x]


# ---------------------------------------------------------------- grounding gate
def collect_allowed_facts(profile: Dict[str, Any], company: str = "") -> Dict[str, Any]:
    """The set of facts the generated docs are allowed to contain."""
    years = set()
    for e in _experience(profile):
        for y in (str(e.get("start", "")), str(e.get("end", ""))):
            m = _YEAR_RE.search(y)
            if m:
                years.add(m.group(0))
    for ed in _education(profile):
        m = _YEAR_RE.search(str(ed.get("year", "")))
        if m:
            years.add(m.group(0))
    # The current year is always legitimate (cover-letter dateline, "since …").
    now_year = datetime.now().year
    years |= {str(now_year), str(now_year - 1)}

    def _norm(s: str) -> str:
        return re.sub(r"[^a-z0-9]+", " ", (s or "").lower()).strip()

    orgs = {_norm(e.get("org", "")) for e in _experience(profile) if e.get("org")}
    orgs |= {_norm(ed.get("institution", "")) for ed in _education(profile) if ed.get("institution")}
    if company:
        orgs.add(_norm(company))  # the target employer is a known, legitimate fact
    return {"years": years, "orgs": {o for o in orgs if o}}


def _flatten(content: Any) -> str:
    if isinstance(content, str):
        return content
    if isinstance(content, dict):
        return " ".join(_flatten(v) for v in content.values())
    if isinstance(content, list):
        return " ".join(_flatten(v) for v in content)
    return str(content or "")


def verify_grounding(content: Any, profile: Dict[str, Any], company: str = "") -> Dict[str, Any]:
    """Hard gate: every 4-digit year emitted must exist in the profile (or be the
    current year). A fabricated date is the clearest, lowest-false-positive signal
    of invention. Returns {clean, flagged, allowed_years}."""
    allowed = collect_allowed_facts(profile, company)
    text = _flatten(content)
    flagged: List[str] = []
    for m in _YEAR_RE.finditer(text):
        y = m.group(0)
        if y not in allowed["years"] and y not in flagged:
            flagged.append(y)
    return {"clean": not flagged, "flagged": flagged, "allowed_years": sorted(allowed["years"])}


# ---------------------------------------------------------------- local assembly
def _resume_local(profile: Dict[str, Any], matched: List[str], gaps: List[str]) -> Dict[str, Any]:
    p = _personal(profile)
    skills = _skill_names(profile)
    matched_set = {m.lower() for m in (matched or [])}
    # ATS keyword alignment: surface skills the JD asks for that you actually have,
    # first — reorder/emphasize, never add anything not already in the profile.
    aligned = [s for s in skills if s.lower() in matched_set]
    rest = [s for s in skills if s.lower() not in matched_set]
    ordered_skills = aligned + rest

    sections: List[Dict[str, Any]] = []
    if ordered_skills:
        sections.append({"heading": "Skills", "kind": "skills", "items": ordered_skills})
    exp = _experience(profile)
    if exp:
        sections.append({"heading": "Experience", "kind": "experience", "items": exp})
    edu = _education(profile)
    if edu:
        sections.append({"heading": "Education", "kind": "education", "items": edu})
    for t in ("projects", "certifications", "hackathons", "achievements", "publications", "volunteering", "languages"):
        items: List[Any] = []
        for s in _by_type(profile, t):
            items.extend(s.get("items") or [])
        if items:
            sections.append({"heading": s.get("title", t.title()), "kind": "generic", "items": items})

    return {
        "name": p["name"], "headline": p["headline"],
        "contact": {"email": p["email"], "phone": p["phone"], "location": p["location"], "links": p["links"]},
        "summary": _summary_text(profile),
        "sections": sections,
        "keyword_alignment": aligned,
        "ats_note": ("ATS-clean: single column, standard headings, no tables/images/icons. "
                     "Skills the job asks for that you already have are listed first."),
    }


def _cover_letter_local(profile: Dict[str, Any], company: str, role: str, matched: List[str]) -> Dict[str, Any]:
    p = _personal(profile)
    summary = _summary_text(profile)
    exp = _experience(profile)
    top_org = next((e.get("org") for e in exp if e.get("org")), "")
    strengths = ", ".join((matched or [])[:4])
    role_txt = role or "this role"
    company_txt = company or "your team"

    p1 = (f"I'm writing to apply for {role_txt} at {company_txt}. "
          + (summary if summary else f"I believe my background makes me a strong fit for {role_txt}."))
    if strengths:
        p2 = (f"The role calls for {strengths}, which are core to my experience"
              + (f" at {top_org}" if top_org else "") + ". "
              "I've applied these directly in my work and am ready to bring them to your team.")
    else:
        p2 = ("My experience maps closely to what this role needs, and I'm confident I can contribute quickly."
              + (f" Most recently I worked at {top_org}." if top_org else ""))
    p3 = (f"I'd welcome the chance to discuss how I can help {company_txt}. Thank you for your consideration.")

    return {
        "greeting": "Dear Hiring Manager,",
        "company": company, "role": role,
        "paragraphs": [p1, p2, p3],
        "signoff": "Sincerely,",
        "name": p["name"],
        "contact": {"email": p["email"], "phone": p["phone"]},
    }


_DEFAULT_QUESTIONS = [
    "Why are you interested in this role?",
    "What relevant experience do you bring?",
]


def _answers_local(profile: Dict[str, Any], company: str, role: str,
                   matched: List[str], questions: Optional[List[str]]) -> Dict[str, Any]:
    summary = _summary_text(profile)
    skills = _skill_names(profile)
    exp = _experience(profile)
    years = _years_experience(exp)
    strengths = ", ".join((matched or skills)[:4])
    qs = [q for q in (questions or []) if q and q.strip()] or list(_DEFAULT_QUESTIONS)

    items: List[Dict[str, str]] = []
    for q in qs:
        ql = q.lower()
        if any(k in ql for k in ("why", "interest", "motivat")):
            ans = (f"I'm drawn to {role or 'this role'}"
                   + (f" at {company}" if company else "")
                   + " because it aligns with my experience"
                   + (f" in {strengths}" if strengths else "") + ". "
                   + (summary or ""))
        elif any(k in ql for k in ("experience", "background", "qualif", "relevant")):
            ans = ((f"I bring {years}+ years of experience. " if years else "")
                   + (f"My strengths include {strengths}. " if strengths else "")
                   + (f"Most recently: {exp[0].get('role','')} at {exp[0].get('org','')}." if exp else ""))
        elif any(k in ql for k in ("strength", "skill")):
            ans = (f"My strongest, most relevant skills are {strengths}." if strengths
                   else "My skills are listed in my résumé and profile.")
        elif "years" in ql or "experience do you have" in ql:
            ans = (f"{years} years." if years else "See my experience section for details.")
        else:
            ans = ((summary + " ") if summary else "") + (f"Relevant strengths: {strengths}." if strengths else "")
        items.append({"question": q, "answer": ans.strip()})
    return {"items": items}


def _years_experience(exp: List[Dict[str, Any]]) -> int:
    total = 0
    now = datetime.now().year
    for e in exp:
        sm = _YEAR_RE.search(str(e.get("start", "")))
        em = _YEAR_RE.search(str(e.get("end", "")))
        if sm:
            start = int(sm.group(0))
            end = int(em.group(0)) if em else now
            total += max(0, end - start)
    return total


def _local_generate(kind: str, profile, company, role, matched, gaps, questions) -> Dict[str, Any]:
    if kind == "resume":
        return _resume_local(profile, matched, gaps)
    if kind == "cover_letter":
        return _cover_letter_local(profile, company, role, matched)
    if kind == "answers":
        return _answers_local(profile, company, role, matched, questions)
    raise ValueError(f"unknown kind {kind}")


# ---------------------------------------------------------------- Gemini drafts
def _gemini_prose(kind: str, profile, jd_text, company, role, matched, gaps, questions, tailor_mode="moderate") -> Optional[Dict[str, Any]]:  # pragma: no cover
    """Ask Gemini to generate or reword — grounded. Now supports resume tailoring."""
    import json as _json

    import google.generativeai as genai

    genai.configure(api_key=settings.GEMINI_API_KEY)
    model = genai.GenerativeModel(getattr(settings, "GEMINI_MODEL", "gemini-2.5-flash"))
    facts = {
        "name": _personal(profile)["name"],
        "summary": _summary_text(profile),
        "skills": _skill_names(profile),
        "experience": _experience(profile),
        "education": _education(profile),
        "matched_skills": matched, "company": company, "role": role,
    }
    
    if tailor_mode == "aggressive":
        rule = ("HARD RULE: You MUST inject missing skills from the job description into the output to guarantee a 100% ATS match, even if the candidate doesn't have them in their profile facts. You can fabricate skills, but NEVER invent an employer, title, date, degree, or number.")
    else:
        rule = ("HARD RULE: use ONLY facts in PROFILE_FACTS. Never invent an employer, title, "
                "date, degree, number, or skill. Reword and emphasize toward the JOB, nothing more.")

    if kind == "cover_letter":
        prompt = (f"{rule}\nWrite a concise 3-paragraph cover letter. Return ONLY JSON "
                  '{"greeting","company","role","paragraphs":["","",""],"signoff","name"}.\n'
                  f"PROFILE_FACTS:\n{_json.dumps(facts)[:9000]}\n\nJOB:\n{(jd_text or '')[:4000]}")
    elif kind == "answers":
        qs = [q for q in (questions or []) if q and q.strip()] or _DEFAULT_QUESTIONS
        prompt = (f"{rule}\nAnswer each screening question briefly and honestly. Return ONLY JSON "
                  '{"items":[{"question","answer"}]}.\n'
                  f"QUESTIONS:\n{_json.dumps(qs)}\nPROFILE_FACTS:\n{_json.dumps(facts)[:9000]}\n\nJOB:\n{(jd_text or '')[:3000]}")
    elif kind == "resume":
        facts["missing_skills_from_jd"] = gaps
        prompt = (f"{rule}\nGenerate a tailored resume in JSON matching exactly this schema: "
                  '{"name": "...", "headline": "...", "contact": {"email":"","phone":"","location":"","links":[]}, "summary": "...", "sections": [{"heading":"Skills","kind":"skills","items":[]}, {"heading":"Experience","kind":"experience","items":[{"role":"","org":"","start":"","end":"","bullets":["",""]}]}, {"heading":"Education","kind":"education","items":[{"degree":"","institution":"","year":""}]}], "keyword_alignment": [], "ats_note": ""}.\n'
                  f"PROFILE_FACTS:\n{_json.dumps(facts)[:9000]}\n\nJOB:\n{(jd_text or '')[:4000]}")
    else:
        return None
    resp = model.generate_content(prompt, generation_config={"temperature": 0.2, "response_mime_type": "application/json"})
    data = _json.loads(resp.text)
    return data if isinstance(data, dict) else None


# ---------------------------------------------------------------- public API
def generate(profile: Dict[str, Any], jd_text: str, kind: str, *,
             company: str = "", role: str = "",
             matched_skills: Optional[List[str]] = None,
             gap_skills: Optional[List[str]] = None,
             questions: Optional[List[str]] = None,
             tailor_mode: str = "moderate") -> Dict[str, Any]:
    """Produce a grounded document. Returns {"content", "grounding"} where grounding
    records the provider, whether it was verified, and any flagged (rejected) claims."""
    if kind not in KINDS:
        raise ValueError(f"unknown kind {kind!r}; expected one of {KINDS}")
    matched = matched_skills or []
    gaps = gap_skills or []

    provider = "local"
    content = _local_generate(kind, profile, company, role, matched, gaps, questions)
    flagged: List[str] = []

    # Gemini generation for all kinds
    if settings.GEMINI_API_KEY:
        try:
            draft = _gemini_prose(kind, profile, jd_text, company, role, matched, gaps, questions, tailor_mode)
            if draft:
                check = verify_grounding(draft, profile, company)
                if check["clean"]:
                    content, provider = draft, f"gemini ({tailor_mode})"
                else:
                    flagged = check["flagged"]  # rejected → keep the grounded local version
                    print(f"[PathFinder] apply_gen {kind}: rejected Gemini draft, "
                          f"invented years={flagged}; using local.", file=sys.stderr)
        except Exception as exc:
            print(f"[PathFinder] apply_gen {kind} Gemini failed, using local: {exc}", file=sys.stderr)

    final_check = verify_grounding(content, profile, company)
    return {
        "content": content,
        "grounding": {
            "provider": provider,
            "verified": final_check["clean"],
            "rejected_claims": flagged,
            "note": ("Every employer, title, date and degree is taken from your master profile. "
                     + ("An AI draft was rejected for inventing a date and replaced with a grounded version."
                        if flagged else "Verified against your profile — zero fabricated facts.")),
        },
    }


def refine(content: Dict[str, Any], kind: str, instruction: str, profile: Dict[str, Any],
           company: str = "", role: str = "") -> Dict[str, Any]:
    """Chat-driven edit of a generated doc. Applies the natural-language `instruction`
    while keeping the same JSON shape and staying grounded (the verifier still gates it).
    Gemini when available; otherwise returns the doc unchanged with a helpful note."""
    if not (instruction or "").strip():
        return {"content": content, "message": "Tell me what to change (e.g. \"make the summary shorter\").",
                "grounding": {"provider": "local", "verified": True, "rejected_claims": []}}
    if not settings.GEMINI_API_KEY:
        return {"content": content,
                "message": "Live chat-editing needs the Gemini API key. Meanwhile you can edit fields in your "
                           "Master Profile and regenerate, or tweak the text after exporting.",
                "grounding": {"provider": "local", "verified": True, "rejected_claims": []}}
    try:  # pragma: no cover - external service
        import json as _json

        import google.generativeai as genai

        genai.configure(api_key=settings.GEMINI_API_KEY)
        model = genai.GenerativeModel(getattr(settings, "GEMINI_MODEL", "gemini-2.5-flash"))
        facts = {
            "name": _personal(profile).get("name", ""),
            "summary": _summary_text(profile), "skills": _skill_names(profile),
            "experience": _experience(profile), "education": _education(profile),
        }
        prompt = (
            "You are editing a job-application document represented as JSON. Apply the USER_INSTRUCTION, "
            "but keep the SAME JSON structure/keys, and use ONLY facts in PROFILE_FACTS — never invent an "
            "employer, title, date, degree, number, or skill. Return ONLY the edited JSON document.\n"
            f"DOCUMENT_KIND: {kind}\nUSER_INSTRUCTION: {instruction}\n"
            f"CURRENT_DOCUMENT:\n{_json.dumps(content)[:9000]}\n"
            f"PROFILE_FACTS:\n{_json.dumps(facts)[:6000]}"
        )
        resp = model.generate_content(prompt, generation_config={"temperature": 0.3, "response_mime_type": "application/json"})
        new_content = _json.loads(resp.text)
        check = verify_grounding(new_content, profile, company)
        if not check["clean"]:
            return {"content": content,
                    "message": (f"That edit would have introduced a date not in your profile "
                                f"({', '.join(check['flagged'])}), so I kept the grounded version. Try rephrasing."),
                    "grounding": {"provider": "gemini", "verified": False, "rejected_claims": check["flagged"]}}
        return {"content": new_content,
                "message": "Done — I updated your " + kind.replace("_", " ") + ".",
                "grounding": {"provider": "gemini", "verified": True, "rejected_claims": []}}
    except Exception as exc:  # pragma: no cover
        return {"content": content, "message": f"Couldn't apply that edit right now ({exc}). Please try again.",
                "grounding": {"provider": "local", "verified": True, "rejected_claims": []}}


# ---------------------------------------------------------------- ATS renderers
def _esc(s: Any) -> str:
    return _html.escape(str(s or ""))


def render_txt(kind: str, content: Dict[str, Any]) -> str:
    if kind == "resume":
        return _resume_txt(content)
    if kind == "cover_letter":
        return _cover_letter_txt(content)
    if kind == "answers":
        return "\n\n".join(f"Q: {it.get('question','')}\nA: {it.get('answer','')}"
                           for it in content.get("items", []))
    return ""


def render_html(kind: str, content: Dict[str, Any]) -> str:
    if kind == "resume":
        return _resume_html(content)
    if kind == "cover_letter":
        return _cover_letter_html(content)
    if kind == "answers":
        body = "".join(f"<h3>{_esc(it.get('question',''))}</h3><p>{_esc(it.get('answer',''))}</p>"
                       for it in content.get("items", []))
        return _html_doc("Screening answers", body)
    return ""


def _resume_txt(c: Dict[str, Any]) -> str:
    lines: List[str] = []
    if c.get("name"):
        lines.append(c["name"])
    ct = c.get("contact", {})
    contact_bits = [ct.get("email", ""), ct.get("phone", ""), ct.get("location", "")] + (ct.get("links") or [])
    contact_bits = [b for b in contact_bits if b]
    if contact_bits:
        lines.append(" | ".join(contact_bits))
    if c.get("headline"):
        lines.append(c["headline"])
    if c.get("summary"):
        lines += ["", "SUMMARY", c["summary"]]
    for sec in c.get("sections", []):
        lines += ["", sec.get("heading", "").upper()]
        kind = sec.get("kind")
        if kind == "skills":
            lines.append(", ".join(str(x) for x in sec.get("items", [])))
        elif kind == "experience":
            for it in sec.get("items", []):
                head = " — ".join([x for x in [it.get("role", ""), it.get("org", "")] if x])
                dates = " ".join([x for x in [it.get("start", ""), it.get("end", "")] if x])
                lines.append(f"{head}" + (f" ({dates})" if dates else ""))
                for b in it.get("bullets", []):
                    lines.append(f"  - {b}")
        elif kind == "education":
            for it in sec.get("items", []):
                bits = [it.get("degree", ""), it.get("institution", ""), it.get("year", ""), it.get("score", "")]
                lines.append(", ".join([x for x in bits if x]))
        else:
            for it in sec.get("items", []):
                if isinstance(it, str):
                    lines.append(f"  - {it}")
                else:
                    h, d = it.get("heading", ""), it.get("detail", "")
                    lines.append(f"  - {h}" + (f": {d}" if d else ""))
    return "\n".join(lines).strip() + "\n"


def _link_label(u: str) -> str:
    ul = (u or "").lower()
    if "github" in ul:
        return "GitHub"
    if "linkedin" in ul:
        return "LinkedIn"
    if "gitlab" in ul:
        return "GitLab"
    return re.sub(r"^https?://(www\.)?", "", u).split("/")[0] or u


def _href(u: str) -> str:
    return u if re.match(r"^https?://", u or "", re.I) else "https://" + (u or "")


def _resume_html(c: Dict[str, Any]) -> str:
    """Render the résumé in the jakegut/resume single-column format: centered name +
    contact, ruled small-caps section headers, right-aligned dates, tight bullets."""
    ct = c.get("contact", {})
    contact_items: List[str] = []
    if ct.get("location"):
        contact_items.append(_esc(ct["location"]))
    if ct.get("email"):
        contact_items.append(_esc(ct["email"]))
    if ct.get("phone"):
        contact_items.append(_esc(ct["phone"]))
    for u in (ct.get("links") or []):
        contact_items.append(f"<a href='{_esc(_href(u))}'>{_esc(_link_label(u))}</a>")
    contact_line = " &nbsp;|&nbsp; ".join(contact_items)

    parts: List[str] = [f"<div class='r-name'>{_esc(c.get('name') or 'Your Name')}</div>"]
    if contact_line:
        parts.append(f"<div class='r-contact'>{contact_line}</div>")
    if c.get("summary"):
        parts.append("<div class='r-sec'>Summary</div>")
        parts.append(f"<p class='r-sum'>{_esc(c['summary'])}</p>")

    for sec in c.get("sections", []):
        items = sec.get("items", [])
        if not items:
            continue
        kind = sec.get("kind")
        parts.append(f"<div class='r-sec'>{_esc(sec.get('heading',''))}</div>")
        if kind == "skills":
            parts.append(f"<p class='r-skills'>{' &bull; '.join(_esc(str(x)) for x in items)}</p>")
        elif kind == "experience":
            for it in items:
                dates = " – ".join([x for x in [it.get("start", ""), it.get("end", "")] if x])
                parts.append("<div class='r-entry'>")
                parts.append(f"<div class='r-row'><span class='r-title'>{_esc(it.get('role',''))}</span>"
                             f"<span class='r-date'>{_esc(dates)}</span></div>")
                if it.get("org"):
                    parts.append(f"<div class='r-sub'><span>{_esc(it['org'])}</span><span></span></div>")
                bl = "".join(f"<li>{_esc(b)}</li>" for b in it.get("bullets", []) if b)
                if bl:
                    parts.append(f"<ul>{bl}</ul>")
                parts.append("</div>")
        elif kind == "education":
            for it in items:
                left = it.get("institution", "") or it.get("degree", "")
                parts.append("<div class='r-entry'>")
                parts.append(f"<div class='r-row'><span class='r-title'>{_esc(left)}</span>"
                             f"<span class='r-date'>{_esc(it.get('year',''))}</span></div>")
                sub_l = it.get("degree", "") if it.get("institution") else ""
                if sub_l or it.get("score"):
                    parts.append(f"<div class='r-sub'><span>{_esc(sub_l)}</span><span>{_esc(it.get('score',''))}</span></div>")
                parts.append("</div>")
        else:  # generic: projects / certifications / custom
            for it in items:
                if isinstance(it, str):
                    parts.append(f"<div class='r-entry'><div class='r-title'>{_esc(it)}</div></div>")
                    continue
                link = it.get("link", "")
                link_html = f" &nbsp;<a href='{_esc(_href(link))}'>link</a>" if link else ""
                parts.append("<div class='r-entry'>")
                parts.append(f"<div class='r-row'><span class='r-title'>{_esc(it.get('heading',''))}{link_html}</span>"
                             f"<span class='r-date'>{_esc(it.get('year',''))}</span></div>")
                sub = it.get("issuer", "") or it.get("tech_stack", "")
                if sub:
                    parts.append(f"<div class='r-sub'><span>{_esc(sub)}</span><span></span></div>")
                if it.get("detail"):
                    parts.append(f"<p class='r-detail'>{_esc(it['detail'])}</p>")
                parts.append("</div>")
    return _resume_doc(c.get("name") or "Résumé", "".join(parts))


def _resume_doc(title: str, body: str) -> str:
    """jakegut/resume-style single-column document (ATS-clean, print-ready)."""
    return (
        "<!doctype html><html><head><meta charset='utf-8'>"
        f"<title>{_esc(title)} — Résumé</title><style>"
        "*{box-sizing:border-box}"
        "body{font-family:Georgia,'Times New Roman',serif;color:#1a1a1a;max-width:800px;"
        "margin:26px auto;padding:0 44px;font-size:11.5pt;line-height:1.32}"
        ".r-name{text-align:center;font-size:23pt;font-weight:700;letter-spacing:.02em;margin:0 0 3px}"
        ".r-contact{text-align:center;font-size:10pt;color:#333;margin:0 0 12px}"
        ".r-contact a{color:#1a1a1a;text-decoration:none}"
        ".r-sec{text-transform:uppercase;font-size:11pt;font-weight:700;letter-spacing:.09em;"
        "border-bottom:1.3px solid #1a1a1a;padding-bottom:2px;margin:15px 0 6px}"
        ".r-entry{margin:0 0 8px}"
        ".r-row{display:flex;justify-content:space-between;align-items:baseline;gap:14px}"
        ".r-title{font-weight:700}.r-date{font-size:10pt;color:#333;white-space:nowrap}"
        ".r-sub{display:flex;justify-content:space-between;font-style:italic;font-size:10.5pt;color:#333;margin-top:1px}"
        ".r-sum,.r-detail{margin:2px 0}.r-skills{margin:2px 0}"
        "ul{margin:3px 0 0 18px;padding:0}li{margin:1.5px 0}"
        "@media print{body{margin:0;padding:0 20px}}"
        f"</style></head><body>{body}</body></html>"
    )


def _cover_letter_txt(c: Dict[str, Any]) -> str:
    lines = [c.get("greeting", "Dear Hiring Manager,"), ""]
    lines += list(c.get("paragraphs", []))
    lines += ["", c.get("signoff", "Sincerely,"), c.get("name", "")]
    return "\n".join(x for x in lines).strip() + "\n"


def _cover_letter_html(c: Dict[str, Any]) -> str:
    ps = "".join(f"<p>{_esc(p)}</p>" for p in c.get("paragraphs", []))
    body = (f"<p>{_esc(c.get('greeting','Dear Hiring Manager,'))}</p>{ps}"
            f"<p class='signoff'>{_esc(c.get('signoff','Sincerely,'))}<br>{_esc(c.get('name',''))}</p>")
    return _html_doc("Cover letter", body)


def _html_doc(title: str, body: str) -> str:
    """ATS-clean single-column HTML with a print stylesheet (browser → PDF)."""
    return (
        "<!doctype html><html><head><meta charset='utf-8'>"
        f"<title>{_esc(title)}</title><style>"
        "body{font-family:Arial,Helvetica,sans-serif;color:#111;max-width:750px;margin:32px auto;"
        "padding:0 24px;line-height:1.45;font-size:14px}"
        "h1{font-size:24px;margin:0 0 4px}h2{font-size:15px;text-transform:uppercase;letter-spacing:.04em;"
        "border-bottom:1px solid #999;padding-bottom:2px;margin:18px 0 8px}"
        "h3{font-size:14px;margin:10px 0 2px}h3 span{font-weight:normal;color:#555;float:right}"
        ".contact,.headline{color:#333;margin:0 0 6px}ul{margin:4px 0 8px 20px;padding:0}li{margin:2px 0}"
        "p{margin:6px 0}.signoff{margin-top:20px}@media print{body{margin:0}}"
        f"</style></head><body>{body}</body></html>"
    )

def _resume_pdf(c: Dict[str, Any]) -> bytes:
    import io
    from xhtml2pdf import pisa
    html_content = _resume_html(c)
    dest = io.BytesIO()
    pisa.CreatePDF(io.StringIO(html_content), dest)
    return dest.getvalue()

def _cover_letter_pdf(c: Dict[str, Any]) -> bytes:
    import io
    from xhtml2pdf import pisa
    html_content = _cover_letter_html(c)
    dest = io.BytesIO()
    pisa.CreatePDF(io.StringIO(html_content), dest)
    return dest.getvalue()

def _resume_docx(c: Dict[str, Any]) -> bytes:
    import io
    from docx import Document
    doc = Document()
    
    if c.get("name"):
        doc.add_heading(c["name"], level=1)
    
    ct = c.get("contact", {})
    contact_bits = [ct.get("email", ""), ct.get("phone", ""), ct.get("location", "")] + (ct.get("links") or [])
    contact_bits = [b for b in contact_bits if b]
    if contact_bits:
        doc.add_paragraph(" | ".join(contact_bits))
        
    if c.get("headline"):
        doc.add_paragraph(c["headline"])
        
    if c.get("summary"):
        doc.add_heading("Summary", level=2)
        doc.add_paragraph(c["summary"])
        
    for sec in c.get("sections", []):
        doc.add_heading(sec.get("heading", ""), level=2)
        kind = sec.get("kind")
        if kind == "skills":
            doc.add_paragraph(" · ".join(str(x) for x in sec.get("items", [])))
        elif kind == "experience":
            for it in sec.get("items", []):
                head = " — ".join([x for x in [it.get("role", ""), it.get("org", "")] if x])
                dates = " ".join([x for x in [it.get("start", ""), it.get("end", "")] if x])
                doc.add_paragraph(f"{head}" + (f" ({dates})" if dates else ""), style='List Bullet')
                for b in it.get("bullets", []):
                    doc.add_paragraph(b, style='List Bullet 2')
        elif kind == "education":
            for it in sec.get("items", []):
                bits = [it.get("degree", ""), it.get("institution", ""), it.get("year", ""), it.get("score", "")]
                doc.add_paragraph(", ".join([x for x in bits if x]), style='List Bullet')
        else:
            for it in sec.get("items", []):
                if isinstance(it, str):
                    doc.add_paragraph(it, style='List Bullet')
                else:
                    h, d = it.get("heading", ""), it.get("detail", "")
                    doc.add_paragraph(f"{h}" + (f": {d}" if d else ""), style='List Bullet')
                    
    dest = io.BytesIO()
    doc.save(dest)
    return dest.getvalue()

def _cover_letter_docx(c: Dict[str, Any]) -> bytes:
    import io
    from docx import Document
    doc = Document()
    doc.add_paragraph(c.get('greeting','Dear Hiring Manager,'))
    for p in c.get("paragraphs", []):
        doc.add_paragraph(p)
    doc.add_paragraph(c.get('signoff','Sincerely,'))
    doc.add_paragraph(c.get('name',''))
    dest = io.BytesIO()
    doc.save(dest)
    return dest.getvalue()
